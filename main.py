import streamlit as st
import openai
from docx import Document
import io
from datetime import datetime
import os
import requests
from docx.shared import RGBColor  # Add this import
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def create_formatted_doc(title, content):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Aptos Narrow'
    
    # Split content into lines
    lines = content.split('\n')
    
    for line in lines:
        # Skip empty lines
        if not line.strip():
            continue
            
        line = line.strip()
        
        # Handle headers
        if line.startswith('# '):
            heading = line.replace('# ', '').strip()
            h = doc.add_heading(heading, level=1)
            h.style.font.color.rgb = RGBColor(0, 51, 102)
        elif line.startswith('## '):
            heading = line.replace('## ', '').strip()
            h = doc.add_heading(heading, level=2)
            h.style.font.color.rgb = RGBColor(0, 51, 102)
        elif line.startswith('### '):
            heading = line.replace('### ', '').strip()
            h = doc.add_heading(heading, level=3)
            h.style.font.color.rgb = RGBColor(0, 51, 102)
        # Handle bullet points
        elif line.startswith('• '):
            # Remove bullet point and handle sub-bullets
            text = line.replace('• ', '').strip()
            indent_level = line.count('\t')  # Count tabs for sub-bullets
            p = doc.add_paragraph(text, style='List Bullet')
            if indent_level > 0:
                p.paragraph_format.left_indent = Pt(indent_level * 18)
        # Regular paragraph
        else:
            doc.add_paragraph(line)
    
    return doc

def process_text(text, detail_level, api_key):
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    prompt = f"""Please analyze these meeting notes and organize them with the following structure:
    # Meeting Title: [Title]
    ## Meeting Notes Summary
    [Brief overview of the meeting]
    
    ## Attendees
    [List attendees with bullet points (•)]
    
    ## Key Points Discussed
    [Main discussion points with bullet points (•)]
    
    Format Guidelines:
    - Use # for main title
    - Use ## for section headers
    - Use • for bullet points
    - Use tabs for sub-bullet points
    - Provide {detail_level} level of detail
    
    Meeting Notes:
    {text}"""
    
    try:
        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json={
                "model": "gpt-4o-mini",
                "messages": [
                    {"role": "system", "content": "You are a professional meeting notes organizer. Format the notes using markdown headers (# for main title, ## for sections) and bullet points (•) for lists."},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.7,
                "max_tokens": 1000
            }
        )
        
        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content']
        else:
            st.error(f"API Error: {response.status_code} - {response.text}")
            return None
            
    except Exception as e:
        st.error(f"Error processing text: {str(e)}")
        return None

def process_text(text, detail_level, api_key):
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    prompt = f"""Please analyze these meeting notes and organize them with the following structure:
    Meeting Title: [Title]
    Meeting Notes Summary
    [Brief overview of the meeting]
    
    Attendees
    [List attendees with bullet points]
    
    Key Points Discussed
    [Main discussion points with bullet points]
    
    Use bullet points (•) for lists
    Use **text** for emphasis/headers within bullet points
    
    Meeting Notes:
    {text}"""
    
    try:
        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json={
                "model": "gpt-4o-mini",
                "messages": [
                    {"role": "system", "content": "You are a professional meeting notes organizer. Format the notes clearly with proper headings and bullet points. Use **text** for emphasis."},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.7,
                "max_tokens": 1000
            }
        )
        
        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content']
        else:
            st.error(f"API Error: {response.status_code} - {response.text}")
            return None
            
    except Exception as e:
        st.error(f"Error processing text: {str(e)}")
        return None

# Streamlit UI
st.title("Meeting Notes Organizer")

st.markdown("""
### Instructions
1. Enter your OpenAI API key
2. Choose between single note input or multiple file upload
3. Select your desired level of detail
4. Process your notes and download the organized version

Created by Brandon Lazovic
""")

# API Key input
api_key = st.text_input("Enter your OpenAI API key", type="password")

# Detail level selector
detail_level = st.select_slider(
    "Select level of detail",
    options=["High-level summary", "Standard detail", "Comprehensive detail"],
    value="Standard detail"
)

# Create tabs for single note and multiple files
tab1, tab2 = st.tabs(["Single Note", "Multiple Files"])

with tab1:
    text_input = st.text_area("Paste your meeting notes here", height=300)
    if st.button("Process Single Note"):
        if not api_key:
            st.error("Please enter an API key")
        elif not text_input:
            st.error("Please enter some text to process")
        else:
            with st.spinner("Processing..."):
                processed_text = process_text(text_input, detail_level, api_key)
                if processed_text:
                    doc = create_formatted_doc("Meeting Notes", processed_text)
                    bio = io.BytesIO()
                    doc.save(bio)
                    st.download_button(
                        label="Download Processed Notes",
                        data=bio.getvalue(),
                        file_name=f"meeting_notes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

with tab2:
    uploaded_files = st.file_uploader(
        "Upload meeting notes (.docx files only)",
        type=['docx'],
        accept_multiple_files=True
    )
    
    if st.button("Process Files") and uploaded_files:
        if not api_key:
            st.error("Please enter an API key")
        else:
            for file in uploaded_files:
                if file.size > 2000000:  # 2MB limit
                    st.error(f"File {file.name} is too large (max 2MB)")
                    continue
                    
                try:
                    doc = Document(file)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    processed_text = process_text(text, detail_level, api_key)
                    
                    if processed_text:
                        new_doc = create_formatted_doc(file.name, processed_text)
                        bio = io.BytesIO()
                        new_doc.save(bio)
                        st.download_button(
                            label=f"Download Processed {file.name}",
                            data=bio.getvalue(),
                            file_name=f"processed_{file.name}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                except Exception as e:
                    st.error(f"Error processing {file.name}: {str(e)}")
